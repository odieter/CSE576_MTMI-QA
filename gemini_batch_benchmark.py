
import os
import json
import time
import google.generativeai as genai
from docx import Document
from tqdm import tqdm


# Load API key from api.txt
with open("api.txt") as f:
    api_key = f.read().strip()

# Configure Gemini
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash-latest")

# Path setup
input_base = "input"
output_base = "output"

# Table formatting function
def format_table_md(table):
    lines = []
    lines.append(f"Table: {table['name']}")
    lines.append("| " + " | ".join(table['columns']) + " |")
    lines.append("|" + "|".join(["---"] * len(table['columns'])) + "|")
    for row in table["data"]:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)

# Process all folders
for person in os.listdir(input_base):
    person_folder = os.path.join(input_base, person)
    if not os.path.isdir(person_folder):
        continue

    files = [f for f in os.listdir(person_folder) if f.endswith(".json")]
    for idx, file in enumerate(files):
        filepath = os.path.join(person_folder, file)
        with open(filepath) as f:
            data = json.load(f)

        formatted_tables = "\n\n".join([format_table_md(t) for t in data["tables"]])
        doc = Document()
        doc.add_heading(f"benchmark_{person.lower()}_{idx+1}", level=1)
        doc.add_paragraph("")

        correct = 0
        partial = 0
        wrong = 0

        for q_idx, qa in tqdm(enumerate(data["questions"]), total=len(data["questions"]), desc=f"Benchmarking {person}_{idx+1}"):
            qid = f"question_fl{str(q_idx+1).zfill(3)}"
        # (rest same)

            prompt = f"""
You are a table reasoning assistant. Use the following tables to answer the question below:

{formatted_tables}

Question: {qa['question']}
Provide a concise answer only.
"""

            response = model.generate_content(prompt)
            model_answer = response.text.strip()

            # Write to DOCX
            doc.add_paragraph(f"Question {qid}: {qa['question']}")
            doc.add_paragraph(f"Gold Answer: {qa['answer']}")
            doc.add_paragraph(f"Model Answer: {model_answer}")

            # Auto check
            gold_answers = [str(ans).lower().strip() for ans in qa["answer"]] if isinstance(qa["answer"], list) else [str(qa["answer"]).lower().strip()]
            model_answer_lower = model_answer.lower().strip()
            model_parts = [part.strip() for part in model_answer_lower.split(",")]

            if set(model_parts) == set(gold_answers):
                doc.add_paragraph("Correct? (Y/N/P): Y\n")
                correct += 1
            elif any(g in part for g in gold_answers for part in model_parts):
                doc.add_paragraph("Correct? (Y/N/P): P\n")
                partial += 1
            else:
                doc.add_paragraph("Correct? (Y/N/P): N\n")
                wrong += 1

            time.sleep(5)

        total = correct + partial + wrong
        doc.add_paragraph(f"Number of Questions: {total}")
        doc.add_paragraph(f"Total Correct: {correct}")
        doc.add_paragraph(f"Total Partially Correct: {partial}")
        doc.add_paragraph(f"Total Incorrect: {wrong}")

        output_path = os.path.join(output_base, f"benchmark_{person.lower()}_{idx+1}.docx")
        doc.save(output_path)
